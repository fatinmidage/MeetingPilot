#!/usr/bin/env python3
"""
会议纪要任务提取工具
从会议记录文本中提取任务信息并生成Excel表格
"""

import os
import sys
import argparse
from pathlib import Path
from typing import List, Dict, Any
from enum import Enum
from datetime import datetime
import glob
from pydantic import BaseModel, Field
from dotenv import load_dotenv
from openai import OpenAI
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document


class TaskType(str, Enum):
    """任务类型枚举"""
    INFO = "信息"
    ACTION = "行动"


class MeetingTask(BaseModel):
    """会议任务数据模型"""
    任务类型: TaskType = Field(description="任务的分类类型：信息（会议中提及的内容）或行动（待办事项）")
    任务描述: str = Field(description="任务的具体描述和要求")
    负责人: str = Field(description="负责执行该任务的人员姓名")
    纳期: str = Field(description="任务的截止日期，格式为YYYY-MM-DD")
    备注: str = Field(description="任务的补充说明或注意事项")


class MeetingResponse(BaseModel):
    """会议响应数据模型"""
    tasks: List[MeetingTask] = Field(description="从会议记录中提取的任务列表")


def load_config() -> dict:
    """加载配置文件"""
    # 从应用程序所在目录加载.env文件
    app_dir = get_application_directory()
    env_path = os.path.join(app_dir, '.env')
    
    # 尝试加载.env文件
    if os.path.exists(env_path):
        load_dotenv(env_path)
    else:
        # 如果应用目录没有.env，尝试当前工作目录
        load_dotenv()
    
    # 获取必需的配置项
    api_key = os.getenv("ARK_API_KEY")
    if not api_key:
        raise ValueError(f"未找到ARK_API_KEY环境变量，请检查.env文件配置\n   期望位置: {env_path}")
    
    model_id = os.getenv("MODEL_ID", "doubao-seed-1.6-250615")
    base_url = os.getenv("BASE_URL", "https://ark.cn-beijing.volces.com/api/v3")
    
    return {
        "api_key": api_key,
        "model_id": model_id,
        "base_url": base_url
    }


def get_application_directory() -> str:
    """获取应用程序所在目录"""
    if getattr(sys, 'frozen', False):
        # 如果是打包后的可执行文件
        return os.path.dirname(sys.executable)
    else:
        # 如果是开发环境运行
        return os.path.dirname(os.path.abspath(__file__))


def discover_meeting_files(directory: str = ".") -> List[Dict[str, Any]]:
    """扫描目录中的会议记录文件"""
    try:
        supported_extensions = ['*.md', '*.txt', '*.docx', '*.doc']
        files = []
        
        # 确定扫描目录：默认为应用程序所在目录
        if directory == ".":
            directory_path = Path(get_application_directory())
        else:
            directory_path = Path(directory)
        
        if not directory_path.exists():
            raise ValueError(f"目录不存在: {directory_path}")
        
        # print(f"[调试] 扫描目录: {directory_path}")  # 可取消注释用于调试
        
        # 扫描所有支持的文件类型
        for pattern in supported_extensions:
            for file_path in directory_path.glob(pattern):
                if file_path.is_file():
                    try:
                        stat = file_path.stat()
                        file_size = stat.st_size
                        
                        # 过滤掉过小的文件
                        if file_size < 10:
                            continue
                        
                        # 格式化文件大小
                        if file_size < 1024:
                            size_display = f"{file_size}B"
                        elif file_size < 1024 * 1024:
                            size_display = f"{file_size / 1024:.1f}KB"
                        else:
                            size_display = f"{file_size / (1024 * 1024):.1f}MB"
                        
                        # 格式化修改时间
                        modified_time = datetime.fromtimestamp(stat.st_mtime)
                        now = datetime.now()
                        
                        if modified_time.date() == now.date():
                            time_display = f"今天 {modified_time.strftime('%H:%M')}"
                        elif (now - modified_time).days == 1:
                            time_display = f"昨天 {modified_time.strftime('%H:%M')}"
                        elif (now - modified_time).days < 7:
                            time_display = f"{(now - modified_time).days}天前"
                        else:
                            time_display = modified_time.strftime('%Y-%m-%d')
                        
                        files.append({
                            "path": str(file_path.resolve()),  # 使用绝对路径确保可执行文件能找到
                            "name": file_path.name,
                            "size": file_size,
                            "size_display": size_display,
                            "modified": modified_time,
                            "time_display": time_display
                        })
                    except (OSError, PermissionError):
                        # 跳过无法访问的文件
                        continue
        
        # 按修改时间倒序排列（最新的在前）
        files.sort(key=lambda x: x["modified"], reverse=True)
        
        return files
        
    except Exception as e:
        raise ValueError(f"扫描文件失败: {e}")


def select_file_interactively(files: List[Dict[str, Any]]) -> str:
    """交互式选择文件"""
    if not files:
        raise ValueError("当前目录未找到支持的会议记录文件")
    
    print("发现以下会议记录文件：")
    for i, file_info in enumerate(files, 1):
        print(f"{i}. {file_info['name']} ({file_info['size_display']}, {file_info['time_display']})")
    
    while True:
        try:
            user_input = input(f"\n请选择要处理的文件 (输入数字 1-{len(files)}，q退出): ").strip()
            
            if user_input.lower() in ['q', 'quit', '退出']:
                print("用户取消操作")
                sys.exit(0)
            
            choice = int(user_input)
            if 1 <= choice <= len(files):
                selected_file = files[choice - 1]
                print(f"已选择: {selected_file['name']}")
                return selected_file['path']
            else:
                print(f"❌ 请输入有效的数字 (1-{len(files)})")
                
        except ValueError:
            print("❌ 请输入有效的数字或 'q' 退出")
        except KeyboardInterrupt:
            print("\n\n用户中断操作")
            sys.exit(0)


def read_docx_file(file_path: str) -> str:
    """读取docx文件内容"""
    try:
        doc = Document(file_path)
        text_content = []
        
        # 提取所有段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 只添加非空段落
                text_content.append(paragraph.text.strip())
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        content = "\n".join(text_content)
        
        if not content:
            raise ValueError(f"Word文档内容为空: {file_path}")
        
        return content
        
    except Exception as e:
        raise ValueError(f"读取Word文档失败: {e}")


def read_meeting_text(file_path: str) -> str:
    """读取会议记录文件，支持多种格式"""
    try:
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        if not file_path_obj.is_file():
            raise ValueError(f"路径不是文件: {file_path}")
        
        # 根据文件扩展名选择读取方式
        file_extension = file_path_obj.suffix.lower()
        
        if file_extension == '.docx':
            # 使用专门的docx读取函数
            content = read_docx_file(file_path)
        elif file_extension == '.doc':
            # .doc格式暂不支持，提示用户转换
            raise ValueError(f"暂不支持.doc格式文件，请将文件转换为.docx格式: {file_path}")
        elif file_extension in ['.md', '.txt', '']:
            # 文本文件，使用UTF-8编码读取
            with open(file_path_obj, 'r', encoding='utf-8') as f:
                content = f.read().strip()
        else:
            raise ValueError(f"不支持的文件格式: {file_extension}，支持的格式: .md, .txt, .docx")
        
        if not content:
            raise ValueError(f"文件内容为空: {file_path}")
        
        return content
        
    except UnicodeDecodeError:
        raise ValueError(f"文件编码错误，请确保文件是UTF-8编码: {file_path}")
    except Exception as e:
        if "不支持" in str(e) or "暂不支持" in str(e):
            raise e
        raise ValueError(f"读取文件失败: {e}")


def extract_tasks(meeting_text: str, config: dict) -> MeetingResponse:
    """使用大模型提取会议任务"""
    try:
        # 初始化OpenAI客户端
        client = OpenAI(
            api_key=config["api_key"],
            base_url=config["base_url"]
        )
        
        # 构建提示词 - 简化为业务逻辑，使用结构化输出
        system_prompt = """你是一个专业的会议纪要分析助手。请仔细分析会议记录内容，提取出所有相关的信息和待办事项。

请将内容分为两种类型：
- 信息：会议中提及的重要内容、决定、讨论结果等信息性内容
- 行动：明确的待办事项、需要执行的任务、后续跟进事项等

对于每项内容，请确定：
- 任务类型：选择"信息"或"行动"
- 任务描述：清晰描述具体内容
- 负责人：相关的责任人姓名，如果没有明确负责人可填"待定"
- 纳期：截止时间或相关时间节点，如果没有明确日期请标注"待定"
- 备注：补充说明、依赖关系或注意事项

请从会议记录中提取所有可识别的信息和行动项。"""

        user_prompt = f"请分析以下会议记录，提取其中的信息和行动项：\n\n{meeting_text}"
        
        # 使用OpenAI结构化输出API
        completion = client.beta.chat.completions.parse(
            model=config["model_id"],
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format=MeetingResponse,  # 指定响应解析模型
            max_tokens=2000,
            temperature=0.1
        )
        
        # 直接获取解析后的结构化响应
        result = completion.choices[0].message.parsed
        
        if not result.tasks:
            raise ValueError("未能从会议记录中提取到信息和行动项")
        
        return result
        
    except Exception as e:
        if "api_key" in str(e).lower() or "unauthorized" in str(e).lower():
            raise ValueError("API密钥验证失败，请检查ARK_API_KEY配置")
        elif "model" in str(e).lower():
            raise ValueError(f"模型调用失败，请检查模型ID配置: {e}")
        else:
            raise ValueError(f"大模型调用失败: {e}")


def generate_excel(tasks: List[MeetingTask], output_path: str) -> None:
    """生成Excel文件"""
    try:
        # 创建工作簿和工作表
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "会议信息行动清单"
        
        # 取消网格线显示
        ws.sheet_view.showGridLines = False
        
        # 定义表头
        headers = ["任务类型", "任务描述", "负责人", "纳期", "备注"]
        
        # 定义边框样式
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # 设置表头样式
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        
        # 写入表头
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border
        
        # 写入任务数据
        for row, task in enumerate(tasks, 2):
            ws.cell(row=row, column=1, value=task.任务类型.value).border = thin_border
            ws.cell(row=row, column=2, value=task.任务描述).border = thin_border
            ws.cell(row=row, column=3, value=task.负责人).border = thin_border
            ws.cell(row=row, column=4, value=task.纳期).border = thin_border
            ws.cell(row=row, column=5, value=task.备注).border = thin_border
        
        # 自动调整列宽
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            
            # 设置列宽（最小10，最大50）
            adjusted_width = min(max(max_length + 2, 10), 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # 设置数据行样式
        data_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = data_alignment
        
        # 保存文件
        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)
        wb.save(output_path)
        
        print(f"Excel文件已生成: {output_path}")
        print(f"共提取项目 {len(tasks)} 项")
        
    except Exception as e:
        raise ValueError(f"生成Excel文件失败: {e}")


def main():
    """主函数入口"""
    # 设置命令行参数解析
    parser = argparse.ArgumentParser(
        description="会议纪要任务提取工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  # 交互模式 (扫描当前目录并选择文件)
  python meeting_extractor.py
  
  # 直接指定文件
  python meeting_extractor.py meeting_notes.md
  python meeting_extractor.py meeting_notes.docx
  python meeting_extractor.py meeting_notes.md -o tasks.xlsx
  
支持的文件格式: .md, .txt, .docx
        """
    )
    
    parser.add_argument(
        "input_file",
        nargs="?",
        help="输入的会议记录文件路径（可选，不提供时将扫描当前目录）"
    )
    
    parser.add_argument(
        "-o", "--output",
        default="meeting_tasks.xlsx",
        help="输出Excel文件路径 (默认: meeting_tasks.xlsx)"
    )
    
    args = parser.parse_args()
    
    try:
        print("=== 会议纪要任务提取工具 ===")
        
        # 确定输入文件
        if args.input_file:
            # 传统模式：用户指定了文件
            input_file = args.input_file
            print(f"输入文件: {input_file}")
        else:
            # 交互模式：扫描目录并让用户选择
            scan_dir = get_application_directory()
            print(f"未指定输入文件，正在扫描程序所在目录: {scan_dir}")
            try:
                available_files = discover_meeting_files()
                if not available_files:
                    print("❌ 程序所在目录未找到支持的会议记录文件")
                    print("   支持的格式: .md, .txt, .docx")
                    print(f"   扫描目录: {scan_dir}")
                    print("   请确保文件存在或使用: meeting_extractor.exe 文件路径")
                    sys.exit(1)
                
                input_file = select_file_interactively(available_files)
            except Exception as e:
                print(f"❌ 文件扫描失败: {e}")
                sys.exit(1)
        
        print(f"输出文件: {args.output}")
        print()
        
        # 1. 加载配置
        print("📋 加载配置...")
        config = load_config()
        print("✅ 配置加载成功")
        
        # 2. 读取会议文本
        print("📖 读取会议记录...")
        meeting_text = read_meeting_text(input_file)
        print(f"✅ 读取成功，文本长度: {len(meeting_text)} 字符")
        
        # 3. 调用大模型提取任务
        print("🤖 调用AI分析会议内容，提取信息和行动项...")
        print("   (这可能需要几秒钟时间)")
        meeting_response = extract_tasks(meeting_text, config)
        print(f"✅ 信息和行动项提取成功，共识别 {len(meeting_response.tasks)} 项")
        
        # 4. 生成Excel文件
        print("📊 生成Excel文件...")
        generate_excel(meeting_response.tasks, args.output)
        print("✅ 处理完成!")
        
        # 5. 显示提取的任务摘要
        print("\n=== 提取摘要 ===")
        for i, task in enumerate(meeting_response.tasks, 1):
            print(f"{i}. {task.任务类型.value} - {task.负责人} - {task.纳期}")
            print(f"   {task.任务描述[:50]}{'...' if len(task.任务描述) > 50 else ''}")
            print()
            
    except KeyboardInterrupt:
        print("\n❌ 用户中断操作")
        sys.exit(1)
    except Exception as e:
        error_msg = str(e)
        if "未找到支持的会议记录文件" in error_msg or "当前目录未找到" in error_msg or "程序所在目录未找到" in error_msg:
            print(f"\n❌ {error_msg}")
            print("💡 建议:")
            print("   1. 将会议记录文件放在可执行文件同一目录下")
            print("   2. 支持的格式: .md, .txt, .docx")
            print("   3. 或直接指定文件: meeting_extractor.exe 文件路径")
        elif "不支持的文件格式" in error_msg or "暂不支持" in error_msg:
            print(f"\n❌ {error_msg}")
            print("💡 建议:")
            print("   1. 使用支持的格式: .md, .txt, .docx")
            print("   2. 如果是.doc文件，请转换为.docx格式")
        elif "Word文档" in error_msg:
            print(f"\n❌ {error_msg}")
            print("💡 建议:")
            print("   1. 确保Word文档没有损坏")
            print("   2. 检查文件权限")
            print("   3. 如果是.doc格式，请转换为.docx")
        elif "ARK_API_KEY" in error_msg:
            print(f"\n❌ {error_msg}")
            print("💡 建议:")
            print("   1. 将 .env 文件复制到可执行文件同一目录下")
            print("   2. 确保 .env 文件包含: ARK_API_KEY=your_actual_api_key_here")
            print("   3. 检查 API 密钥是否正确配置")
        else:
            print(f"\n❌ 错误: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main() 